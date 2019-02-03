from __future__ import absolute_import, unicode_literals
import os, django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "second_Edition.settings")
django.setup()

from . import models
import hashlib
import urllib
import json
import http
import requests
import random
import docx



def fanyi(content):
    appid = '20181215000248881'
    secretKey = '6V2UsBsIaP7vGqwknI4K'
    httpClient = None
    myurl = '/api/trans/vip/translate'
    q = content
    fromLang = 'zh'
    toLang = 'en'
    salt = random.randint(32768, 65536)
    sign = appid + q + str(salt) + secretKey
    sign = hashlib.md5(sign.encode()).hexdigest()
    myurl = myurl + '?appid=' + appid + '&q=' + urllib.parse.quote(
        q) + '&from=' + fromLang + '&to=' + toLang + '&salt=' + str(
        salt) + '&sign=' + sign

    try:
        httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
        httpClient.request('GET', myurl)
        response = httpClient.getresponse()
        jsonResponse = response.read().decode("utf-8")
        js = json.loads(jsonResponse)
        dst = str(js["trans_result"][0]["dst"])
        print(dst)
        return dst
    except Exception as e:
        print(e)
        return e
    finally:
        if httpClient:
            httpClient.close()


def enfanyi(content):
    appid = '20181215000248881'
    secretKey = '6V2UsBsIaP7vGqwknI4K'
    httpClient = None
    myurl = '/api/trans/vip/translate'
    q = content
    fromLang = 'en'
    toLang = 'zh'
    salt = random.randint(32768, 65536)
    sign = appid + q + str(salt) + secretKey
    sign = hashlib.md5(sign.encode()).hexdigest()
    myurl = myurl + '?appid=' + appid + '&q=' + urllib.parse.quote(
        q) + '&from=' + fromLang + '&to=' + toLang + '&salt=' + str(
        salt) + '&sign=' + sign

    try:
        httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
        httpClient.request('GET', myurl)
        response = httpClient.getresponse()
        jsonResponse = response.read().decode("utf-8")
        js = json.loads(jsonResponse)
        dst = str(js["trans_result"][0]["dst"])
        print(dst)
        return dst
    except Exception as e:
        print(e)
        return e
    finally:
        if httpClient:
            httpClient.close()


class Sogou(object):
    def __init__(self):
        # PID
        self.pid = 'd2dcca8dea30c3dc0570788c204f5b8a'
        # KEY
        self.key = '475af5d76c46584ec5d35493b92f255b'
        # url
        self.url = 'http://fanyi.sogou.com:80/reventondc/api/sogouTranslate'
        # headers
        self.headers = {
            'content-type': "application/x-www-form-urlencoded",
            'accept': "application/json"
        }
    def en_zh(self, q, salt):
        sign = self.pid + q + salt + self.key
        sign = hashlib.md5(sign.encode('utf8')).hexdigest()
        payload = {'pid': self.pid, 'q': q.encode('utf8'), 'salt': salt, 'sign': sign}
        print(payload)
        response = requests.post(self.url, data=payload, headers=self.headers)
        print(response.text)
        enres = json.loads(response.text)
        enres = enres['translation']
        return enres

    def zh_en(self, q, salt):
        sign = self.pid + q + salt + self.key
        sign = hashlib.md5(sign.encode('utf8')).hexdigest()
        payload = {'pid': self.pid, 'q': q.encode('utf8'), 'salt': salt, 'sign': sign}
        # print(payload)
        response = requests.post(self.url, data=payload, headers=self.headers)
        print(response.text)
        res = json.loads(response.text)
        res = res['translation']
        return res


en = Sogou()


# @celery.app.task
def zh_rs(path, user, update_file, zl):
    file = docx.Document(path)
    zl = int(zl)
    print("翻译中->Ying类", zl)
    print(type(zl))
    for p in file.paragraphs:
        for r in p.runs:
            if r.bold:
                print(r.bold)
                print( r.text)
                fanyi_save(zl, user, r.text, 'bold', r.bold, update_file)
            elif r.font.color.rgb:
                print(r.font.color.rgb)
                print(r.text)
                fanyi_save(zl, user, r.text, 'font.color.rgb', r.font.color.rgb, update_file)
            elif r.font.italic:
                print(r.font.italic)
                print(r.text)
                fanyi_save(zl, user, r.text, 'font.italic', r.font.italic, update_file)
            elif r.font.underline:
                print(r.font.underline)
                print(r.text)
                fanyi_save(zl, user, r.text, 'font.underline', r.font.underline, update_file)
            else:
                print(r.text)
                fanyi_save(zl, user, r.text, p.style.name, r.style.name, update_file)


def fanyi_save(zl, user, content, content_style, value, update_file):
    if zl == 1:
        if len(content) > 1:
            yiwen = fanyi(content)
        else:
            yiwen = ""
    if zl == 2:
        if len(content.strip()) > 1:
            # time.sleep(1)
            # en = Sogou()
            print(content)
            pp = content.strip()
            yiwen = en.zh_en(pp, str(random.randint(1000000000000, 9999999999999)))
            # yiwen = sougou.sougou_zh(p.text)
        else:
            yiwen = ""
    print("*************")
    print(yiwen)
    print("*************")
    # 一次传入多少字节
    once_num = len(content)
    line = models.yuanwen(yuanwen=content, yuanwen_style=content_style, style_value = value, user=user, file=update_file)
    line.save()
    # 将每次的字节数量+=到数据库中
    file = models.file_information.objects.get(pk=update_file.id)
    file.all_num += once_num
    # 传入一次即为加一段
    file.all_para += 1
    file.save()
    yw = models.yiwen(yiwen="", yiwen_wait_over=yiwen, yiwen_style=content_style, gl=line, user=user, style_value = value, file=update_file)
    yw.save()


# @celery.app.task
def en_read_save(path, user, update_file, zl):
    file = docx.Document(path)
    zl = int(zl)
    print("英-》中", zl)
    print(type(zl))
    for p in file.paragraphs:
        for r in p.runs:
            if r.bold:
                print(r.bold)
                en_fanyi_save(zl, user, r.text, 'bold', r.bold, update_file)
            elif r.font.color.rgb:
                print(r.font.color.rgb)
                en_fanyi_save(zl, user, r.text, 'font.color.rgb', r.font.color.rgb, update_file)
            elif r.font.italic:
                print(r.font.italic)
                en_fanyi_save(zl, user, r.text, 'font.italic', r.font.italic, update_file)
            elif r.font.underline:
                print(r.font.underline)
                en_fanyi_save(zl, user, r.text, 'font.underline', r.font.underline, update_file)
            else:
                print(r.text)
                en_fanyi_save(zl, user, r.text, p.style.name, r.style.name, update_file)


def en_fanyi_save(zl, user, content, content_style, value, update_file):
    if zl == 1:
        if len(content) > 1:
            yiwen = enfanyi(content)
        else:
            yiwen = ""
    if zl == 2:
        if len(content.strip()) > 1:
            # time.sleep(1)
            # en = Sogou()
            print(content)
            pp = content.strip()
            yiwen = en.en_zh(pp, str(random.randint(1000000000000, 9999999999999)))
            # yiwen = sougou.sougou_zh(p.text)
        else:
            yiwen = ""
    print("*************")
    print(yiwen)
    print("*************")
    # 一次传入多少字节
    once_num = len(content)
    line = models.yuanwen(yuanwen=content, yuanwen_style=content_style, style_value = value, user=user, file=update_file)
    line.save()
    # 将每次的字节数量+=到数据库中
    file = models.file_information.objects.get(pk=update_file.id)
    file.all_num += once_num
    # 传入一次即为加一段
    file.all_para += 1
    file.save()
    yw = models.yiwen(yiwen="", yiwen_wait_over=yiwen, yiwen_style=content_style, gl=line, user=user, style_value = value, file=update_file)
    yw.save()
