from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, StreamingHttpResponse, response
# 登录检查
from django.contrib.auth.decorators import login_required
# 限制访问方式
from django.views.decorators.http import require_http_methods
# 引入数据库
from . import models
# 引入事务
from django.db import transaction
# 引入用户
from . import task
import docx
from .task import fanyi, enfanyi, fanyi_save, en_fanyi_save, Sogou
import random
import subprocess



# 首页面
@login_required
def index(request):
    files = models.file_information.objects.filter(file_user=request.user)
    return render(request, 'work/index.html', {"files":files})


# 上传首页面
@login_required
def update(request):
    return render(request, 'work/update.html', {})


# 文件上传
@login_required
@require_http_methods(["POST"])
def file_update(request):
    # try:
    print("开始获取")
    # 获取文件和名称
    file = request.FILES.get('file')
    name = file.name
    # 获取当前的用户
    user = request.user
    print("获取到，存储")
    print(file, name, user)
    # 存入数据库
    update_file = models.file_information(file_name = name, file=file, file_user=user)
    print("???")
    print(update_file)
    update_file.save()
    # 根据文件位置进行修改
    doc_path = "/home/admin/second_Edition/" + str(update_file.file)
    print(doc_path)
    id = update_file.id
    if doc_path.split('.', -1)[-1] == "doc":
        print("不是docx")
        a = subprocess.check_output(["soffice", "--headless",
                                     "--invisible", "--convert-to",
                                     "docx", doc_path,
                                     "--outdir", "/home/admin/second_Edition/static/work/doc"])
        name = str(update_file.file) + "x"
        print("docx" + name)
        print(a)
        print(name)
        docx_save = models.file_information.objects.get(pk=id)
        docx_save.file = name
        docx_save.save()
    print("end")

    return render(request, 'work/file_info.html', {"msg": file.name, 'id': update_file.id})
    # except:
    #     return render(request, 'work/file_info.html', {"msg": "cuowu "})


# 文件信息填写
@login_required
@require_http_methods(["POST"])
def file_info(request):
    name = request.POST['name']
    info = request.POST['info']
    over_time = request.POST['over_time']
    id = request.POST['id']     # 文件ID
    print(id)
    yuan = request.POST['yuan']
    print("源："+yuan)
    mubiao = request.POST['mubiao']
    print("目标："+ mubiao)
    print("时间", over_time)
    # 查询当前的数据
    file = models.file_information.objects.get(pk=id)
    file.file_name = name
    file.file_info = info
    file.yuan = yuan
    file.mubiao = mubiao
    file.over_time = over_time
    file.save()
    print("保存完成")
    return render(request, 'work/trans_type.html', {'id': id})


# 选择翻译种类
@login_required
@require_http_methods(["POST"])
def trans_type(request):
    id = request.POST['id']
    zl = request.POST['zl']
    print("翻译的种类"+zl)
    file = models.file_information.objects.get(pk=id)
    file.fanyi_zl= zl
    file.save()
    print(file)
    fx = file.mubiao
    print("fx", fx)
    # zl = file.fanyi_zl
    # print("zl:",zl)
    if fx == 2:
        print("等于1")
        return render(request, 'work/work_wait.html', {"file":file})
    if fx == 1:
        print("等于2")
        return render(request, 'work/work_enwait.html', {"file":file})


@login_required
# @require_http_methods(["POST"])
# def work_wait(request, id, fx, yuan):
def work_wait(request):
    return render(request, 'work/work_wait.html', {})


# 中文编译界面
@login_required
def work(request, file_id):
    file = models.file_information.objects.get(pk = file_id)
    print(file.file_name)
    # 获取当前的用户
    user = request.user
    # 获取当前ID的用户
    id_user = file.file_user
    # 判断是否用户是否为一人
    if user == id_user:
        yuanwen = models.yuanwen.objects.filter(file = file_id)
        print(len(yuanwen))
        print(yuanwen)
        if len(yuanwen) > 1:
            yiwen = models.yiwen.objects.filter(file = file_id)
            file = models.file_information.objects.get(pk = file_id)
            wenjian = zip(yuanwen, yiwen)
            return render(request, 'work/work.html', {"wenjian": wenjian, "file_id": file_id, "filename":file.file_name, "file":file})
        else:
            path = file.file
            path = str(path)
            task.zh_rs(path, request.user, file, file.mubiao)
            return redirect("/work/" + file_id + "/work/")
    else:
        return render(request, 'work/work_error.html', {})


# 英文编译界面
@login_required
def enwork(request, file_id):
    file = models.file_information.objects.get(pk=file_id)
    print(file.file_name)
    # 获取当前的用户
    user = request.user
    # 获取当前ID的用户
    id_user = file.file_user
    # 判断是否用户是否为一人
    if user == id_user:
        yuanwen = models.yuanwen.objects.filter(file=file_id)
        print(len(yuanwen))
        print(yuanwen)
        if len(yuanwen) > 1:
            yiwen = models.yiwen.objects.filter(file=file_id)
            file = models.file_information.objects.get(pk=file_id)
            wenjian = zip(yuanwen, yiwen)
            return render(request, 'work/enwork.html',
                          {"wenjian": wenjian, "file_id": file_id, "filename": file.file_name, "file": file})
        else:
            path = file.file
            path = str(path)
            task.en_read_save(path, request.user, file, file.mubiao)
            return redirect("/work/" + file_id + "/work/")
    else:
        return render(request, 'work/work_error.html', {})


import re
# 将获取到的16进制颜色转化为rgb颜色
def toRgb(tmp):
    opt = re.findall(r'(.{2})', tmp)  # 将字符串两两分割
    strs = ""  # 用以存放最后结果
    for i in range(0, len(opt)):  # for循环，遍历分割后的字符串列表
        strs += str(int(opt[i], 16)) + ","  # 将结果拼接成12，12，12格式
    return strs[0:-1]


from docx.shared import RGBColor


def word_ok(file_id):
    file = models.file_information.objects.get(pk = file_id)
    print("word名称：", file.file_name)
    document = docx.Document()
    yiwen = models.yiwen.objects.filter(file = file)
    for y in yiwen:
        print("1"+y.yiwen_wait_over)
        print("2"+y.yiwen_style)
        print("3"+y.style_value)
        one_list = ['Normal', 'List Paragraph', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Title', 'Subtitle']
        if len(y.yiwen_style) > 0 and y.yiwen_style in one_list:
            p = document.add_paragraph(y.yiwen_wait_over, style=y.yiwen_style)
            p.alignment = 3
            # p.alignment = alignment

    path = "static/work/yiwen_doc/" + file.file_name
    # 把生成的word文档路径存到数据库
    sqL_oath_save = models.file_information.objects.get(id = file_id)
    sqL_oath_save.word_path = path
    sqL_oath_save.save()
    document.save(path)
    return path, file.file_name


# 缓冲流下载文件方法
def readFile(filename, chunk_size=512):
    with open(filename, 'rb') as f:
        while True:
            c = f.read(chunk_size)
            if c:
                yield c
            else:
                break


# 导出word
@login_required
# @require_http_methods(["POST"])
def downloads(request, file_id):
    file = models.file_information.objects.get(pk = file_id)
    print(file.file_name)
    # 获取当前的用户
    user = request.user
    # 获取当前ID的用户
    id_user = file.file_user
    # 判断是否用户是否为一人
    if user == id_user:
        path, file_name = word_ok(file_id)
        print(path)
        print(file_name)
        response = StreamingHttpResponse(readFile(path))
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = 'attachment;filename="download.docx"'
        # 把状态改为已完成
        file = models.file_information.objects.get(pk = file_id)
        file.file_status = 0
        file.save()
        return response
    else:
        return render(request, 'work/update.html', {})


# 删除文件
@login_required
# @require_http_methods(["POST"])
def del_file(request, file_id):
    file = models.file_information.objects.get(pk = file_id)
    print(file.file_name)
    # 获取当前的用户
    user = request.user
    # 获取当前ID的用户
    id_user = file.file_user
    # 判断是否用户是否为一人
    if user == id_user:
        file = models.file_information.objects.get(pk = file_id)
        file.file_status = 2
        file.save()
        return redirect('/work/')
    else:
        return render(request, 'work/work_error.html', {})

import json
from django.http import HttpResponse
# 单行/全部提交译文
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def change_yiwen(request, id, yiwen):
    print(yiwen, id)
    # 查找译文表，把提交的译文存入到译文列和待修改列
    p = models.yiwen.objects.get(gl_id = id)
    p.yiwen = yiwen
    p.yiwen_wait_over = yiwen
    p.save()
    # 查找原文表，得出原文的长度
    yuanwen = models.yuanwen.objects.get(pk=id)
    yuanwen_len = len(yuanwen.yuanwen)
    print("原文长：", yuanwen_len)
    print(yuanwen.file.id)
    # 查找文件信息表，把提交来的译文的原文长度存到数据库中 并将是否可修改改为1（不可修改）
    file = models.file_information.objects.get(pk=yuanwen.file.id)
    if yuanwen.check_changeed != 1:
        yuanwen.check_changeed = 1
        file.numing += yuanwen_len
        file.paraing += 1
        jindu = int((yuanwen_len+file.numing)/file.all_num*100)
        if jindu >= 100:
            jindu = 100
        file.jindu = jindu
        file.save()
        yuanwen.save()
    resp = {'numing': file.numing, 'paraing': file.paraing, 'jindu':file.jindu, "all_num":file.all_num, "all_para":file.all_para}
    return HttpResponse(json.dumps(resp), content_type="application/json")


# 全部保存译文
@login_required
@require_http_methods(["POST"])
def all_save_yiwen(request, y_id, yiwen):
    try:
        print(y_id, yiwen)
        p = models.yiwen.objects.get(gl_id=y_id)
        print(p)
        if request.method == 'POST':
            print(yiwen)
            p.yiwen_wait_over = yiwen
            p.save()
            return JsonResponse({"msg": "全部保存成功", "success": True})
    except:
        return JsonResponse({"msg": "全部保存失败", "success": False})


# 搜索原文
@login_required
@require_http_methods(["POST"])
def serch_yuanwen(request, file_id):
    if request.method == 'POST':
        yuanwen_keyword = request.POST['yuanwen_keyword']
        yuanwen_list = []
        yiwen_list = []
        yuanwen = models.yuanwen.objects.filter(file = file_id)
        for y in yuanwen:
            # print(y.yuanwen)
            if yuanwen_keyword in y.yuanwen:
                print("找到了：", y.yuanwen, y.id)
                yiwen_id = y.id
                yuanwen_list.append(y)
                yiwen = models.yiwen.objects.get(gl = yiwen_id)
                yiwen_list.append(yiwen)
                print(yiwen.yiwen)
        print(yuanwen_list)
        print(yiwen_list)
        wenjian = zip(yuanwen_list, yiwen_list)
        return render(request, 'work/serch.html',
                      {"wenjian": wenjian, "file_id": file_id})


# 搜索译文
@login_required
@require_http_methods(["POST"])
def serch_yiwen(request, file_id):
    if request.method == 'POST':
        yiwen_keyword = request.POST['yiwen_keyword']
        yuanwen_list = []
        yiwen_list = []
        print("接收到：",yiwen_keyword)
        yiwen = models.yiwen.objects.filter(file = file_id)
        for y in yiwen:
            print(y.yiwen_wait_over)
            if yiwen_keyword in y.yiwen_wait_over:
                print("找到了：", y.yiwen_wait_over, y.gl_id)
                yuanwen_id = y.gl_id
                yiwen_list.append(y)
                yuanwen = models.yuanwen.objects.get(pk = yuanwen_id)
                print("原文",yuanwen)
                # yuanwen = yuanwen
                yuanwen_list.append(yuanwen)
                # print(yiwen.yiwen)
        print(yuanwen_list)
        print(yiwen_list)
        wenjian = zip(yuanwen_list, yiwen_list)
        return render(request, 'work/serch.html',
                      {"wenjian": wenjian, "file_id": file_id})


# 中文
# json百度翻译
@login_required
@require_http_methods(["POST"])
def bdjson(request, query, id):
    print("1ojbk")
    if request.method == 'POST':
        if len(query) > 1:
            yiwen = fanyi(query)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 搜狗翻译
@login_required
@require_http_methods(["POST"])
def sgjson(request, query, id):
    print("ojbk")
    if request.method == 'POST':
        if len(query) > 1:
            en = Sogou()
            query = query.strip()
            yiwen = en.zh_en(query, str(random.randint(1000000000000, 9999999999999)))
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 英文
# json百度翻译
@login_required
@require_http_methods(["POST"])
def en_baidu(request, query, id):
    print("enojbk")
    if request.method == 'POST':
        if len(query) > 1:
            print("原"+query)
            yiwen = enfanyi(query)
            print(enfanyi(query))
            print("yi"+yiwen)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 搜狗翻译
@login_required
@require_http_methods(["POST"])
def ensgjson(request, query, id):
    print("ojbk")
    if request.method == 'POST':
        if len(query) > 1:
            en = Sogou()
            yiwen = en.en_zh(query, str(random.randint(1000000000000, 9999999999999)))
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 查看修改前
def change_befor(request, id):
    yiwen = models.yiwen.objects.get(pk=id)
    print(yiwen.yiwen)
    print(yiwen.gl.yuanwen)
    listlist = [yiwen.yiwen_wait_over, yiwen.gl.yuanwen]
    return JsonResponse({"msg": listlist, "success": True})


# 查看修改后
def change_old(request, id):
    yiwen = models.yiwen.objects.get(pk=id)
    print(yiwen.yiwen_change)
    print(yiwen.gl.yuanwen_change)
    listlist = [yiwen.yiwen_change, yiwen.gl.yuanwen_change]
    return JsonResponse({"msg": listlist, "success": True})


# 编译原文栏后的自动机器翻译
@login_required
@require_http_methods(["POST"])
def change_content(request, query, id):
    print("oo")
    yiwen_change = fanyi(query)
    print("译文："+yiwen_change)
    yuanwen_change = query
    print("原文："+yuanwen_change)
    yiwen = models.yiwen.objects.get(gl_id=id)
    yuanwen = models.yuanwen.objects.get(pk=id)
    yiwen.yiwen_change = yiwen_change
    print('y1')
    yuanwen.yuanwen_change= yuanwen_change
    print("y2")
    yiwen.save()
    yuanwen.save()
    return JsonResponse({"msg": yiwen_change, "success": True})


# 英文 ：编译原文栏后的自动机器翻译
@login_required
@require_http_methods(["POST"])
def en_content_change(request, query, id):
    print("enoo")
    yiwen_change = enfanyi(query)
    print("译文："+yiwen_change)
    yuanwen_change = query
    print("原文："+yuanwen_change)
    yiwen = models.yiwen.objects.get(gl_id=id)
    yuanwen = models.yuanwen.objects.get(pk=id)
    yiwen.yiwen_change = yiwen_change
    print('y1')
    yuanwen.yuanwen_change= yuanwen_change
    print("y2")
    yiwen.save()
    yuanwen.save()
    return JsonResponse({"msg": yiwen_change, "success": True})


# 检查是否修改了原文栏
def yuanwen_check(request, id):
    yuanwen = models.yuanwen.objects.get(pk=id)
    print(yuanwen.yuanwen)
    return JsonResponse({"msg": yuanwen.yuanwen, "success": True})


# 检查是否修改了译文
def yiwen_check(request, id):
    yiwen = models.yiwen.objects.get(gl=id)
    print(yiwen.yiwen_wait_over)
    return JsonResponse({"msg": yiwen.yiwen_wait_over, "success": True})


def zh_rs(request, id):
    print("进来了")
    print("文件ID", id)
    word_file = models.file_information.objects.get(pk=id)
    path = word_file.file
    zl = word_file.fanyi_zl
    file = docx.Document(path)
    zl = int(zl)
    print("翻译中->Ying类", zl)
    print(type(zl))
    user = request.user
    print("yonhu:", user)
    for p in file.paragraphs:
        print(p.text)
        print(p.style.name)
        if zl == 1:
            print("使用百度")
            # 预翻译（百度）
            if len(p.text) > 1:
                yiwen = fanyi(p.text)
            else:
                yiwen = ""
        if zl == 2:
            print("使用搜狗")
            if len(p.text.strip()) > 1:
                # time.sleep(1)
                # en = Sogou()
                print(p.text)
                pp = p.text.strip()
                yiwen = task.en.en_zh(pp, str(random.randint(1000000000000, 9999999999999)))
                # yiwen = sougou.sougou_zh(p.text)
            else:
                yiwen = ""
        print("*************")
        print(yiwen)
        print("*************")
        # 带标题的原文存入数据库
        line = models.yuanwen(yuanwen=p.text, yuanwen_style=p.style.name, user=user, file=word_file)
        line.save()
        # 译文存入数据库
        yw = models.yiwen(yiwen_wait_over=yiwen, yiwen_style=p.style.name, gl=line, user=user, file=word_file)
        yw.save()
    return JsonResponse({"msg": "OK", "success": True})


def en_read_save(request, id):
    print("进来了")
    print("文件ID", id)
    word_file = models.file_information.objects.get(pk=id)
    path = word_file.file
    zl = word_file.fanyi_zl
    file = docx.Document(path)
    zl = int(zl)
    print("翻译中->Ying类", zl)
    print(type(zl))
    user = request.user
    print("yonhu:", user)
    for p in file.paragraphs:
        print(p.text)
        print(p.style.name)
        if zl == 1:
            print("使用百度")
            # 预翻译（百度）
            if len(p.text) > 1:
                yiwen = enfanyi(p.text)
            else:
                yiwen = ""
        if zl == 2:
            print("使用搜狗")
            if len(p.text) > 1:
                # import time
                # time.sleep(1)
                # en = Sogou()
                yiwen = task.en.en_zh(p.text, str(random.randint(1000000000000, 9999999999999)))
            else:
                yiwen = ""
        print("*************")
        print(yiwen)
        print("*************")
        # 带标题的原文存入数据库
        line = models.yuanwen(yuanwen=p.text, yuanwen_style=p.style.name, user=user, file=word_file)
        line.save()
        # 译文存入数据库
        yw = models.yiwen(yiwen_wait_over=yiwen, yiwen_style=p.style.name, gl=line, user=user, file=word_file)
        yw.save()

    return JsonResponse({"msg": "OK", "success": True})

